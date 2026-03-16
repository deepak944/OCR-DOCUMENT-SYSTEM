from collections import defaultdict
from io import BytesIO
import logging

import fitz
from docx import Document
from docx.shared import Inches
from PIL import Image

from app.services.document_services import process_document

WORD_IMAGE_WIDTH_INCHES = 5.8


def _safe_sort_key(value):
    try:
        return (0, int(value))
    except (TypeError, ValueError):
        return (1, str(value))


def _normalize_text(value):
    lines = [line.strip() for line in str(value or "").splitlines() if line.strip()]
    return " ".join(lines)


def _bbox_sort_key(block):
    bbox = block.get("bbox") if isinstance(block, dict) else None

    if isinstance(bbox, list):
        if bbox and isinstance(bbox[0], (list, tuple)) and len(bbox[0]) >= 2:
            try:
                x, y = bbox[0][0], bbox[0][1]
                return (float(y), float(x))
            except (TypeError, ValueError):
                pass

        if len(bbox) >= 2:
            try:
                x, y = bbox[0], bbox[1]
                return (float(y), float(x))
            except (TypeError, ValueError):
                pass

    return (float("inf"), float("inf"))


def _table_dict_to_rows(table_dict):
    if not isinstance(table_dict, dict):
        return []

    columns = []
    for column_key, row_map in table_dict.items():
        if isinstance(row_map, dict):
            columns.append((column_key, row_map))

    if not columns:
        return []

    ordered_columns = [col for col, _ in sorted(columns, key=lambda item: _safe_sort_key(item[0]))]
    row_keys = set()
    for _, row_map in columns:
        row_keys.update(row_map.keys())

    ordered_rows = sorted(row_keys, key=_safe_sort_key)
    rows = []

    for row_key in ordered_rows:
        row_values = []
        for column_key in ordered_columns:
            raw_value = table_dict.get(column_key, {}).get(row_key, "")
            row_values.append(_normalize_text(raw_value))
        rows.append(row_values)

    return rows


def _append_tables(document, tables):
    table_count = 0

    for table_data in tables:
        rows = _table_dict_to_rows(table_data)
        if not rows:
            continue

        if table_count == 0:
            document.add_heading("Extracted Tables", level=1)

        table_count += 1
        document.add_paragraph(f"Table {table_count}")

        doc_table = document.add_table(rows=len(rows), cols=len(rows[0]))
        doc_table.style = "Table Grid"

        for row_index, row in enumerate(rows):
            for col_index, value in enumerate(row):
                doc_table.cell(row_index, col_index).text = value


def _normalized_page_number(value):
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _group_images_by_page(images):
    grouped_images = defaultdict(list)

    for image in images:
        if not isinstance(image, dict):
            continue

        page_number = _normalized_page_number(image.get("page_number"))
        if page_number is None:
            continue

        grouped_images[page_number].append(image)

    return grouped_images


def _image_sort_key(image):
    if not isinstance(image, dict):
        return (float("inf"), float("inf"))

    image_index = image.get("image_index")
    xref = image.get("xref")
    return (_safe_sort_key(image_index), _safe_sort_key(xref))


def _extract_image_bytes(pdf_document, image_meta, page_number):
    xref = _normalized_page_number(image_meta.get("xref") if isinstance(image_meta, dict) else None)

    if xref is None:
        return None

    try:
        image_data = pdf_document.extract_image(xref)
    except Exception:
        logging.exception("Unable to extract image xref=%s for page %s", xref, page_number)
        return None

    return image_data.get("image")


def _to_docx_compatible_image_stream(raw_image_bytes):
    if not raw_image_bytes:
        return None

    try:
        with Image.open(BytesIO(raw_image_bytes)) as image:
            prepared_image = image
            if prepared_image.mode in ("CMYK", "P"):
                prepared_image = prepared_image.convert("RGB")

            output_stream = BytesIO()
            prepared_image.save(output_stream, format="PNG")
            output_stream.seek(0)
            return output_stream
    except Exception:
        fallback_stream = BytesIO(raw_image_bytes)
        fallback_stream.seek(0)
        return fallback_stream


def _append_page_images(document, pdf_document, page_number, page_images):
    if not page_images:
        return

    document.add_paragraph("Extracted images:")
    embedded_count = 0

    for image_position, image_meta in enumerate(sorted(page_images, key=_image_sort_key), start=1):
        raw_image_bytes = _extract_image_bytes(pdf_document, image_meta, page_number)
        image_stream = _to_docx_compatible_image_stream(raw_image_bytes)

        if image_stream is None:
            continue

        document.add_paragraph(f"Image {image_position}")

        try:
            image_paragraph = document.add_paragraph()
            image_paragraph.add_run().add_picture(image_stream, width=Inches(WORD_IMAGE_WIDTH_INCHES))
            embedded_count += 1
        except Exception:
            logging.exception("Failed to embed image %s on page %s", image_position, page_number)

    if embedded_count == 0:
        document.add_paragraph("Images were detected on this page, but they could not be embedded.")


def convert_pdf_to_word_doc(pdf_path, output_doc_path):
    extracted = process_document(pdf_path)
    pages = extracted.get("pages", [])
    tables = extracted.get("tables", [])
    images = extracted.get("images", [])
    images_by_page = _group_images_by_page(images)

    document = Document()
    document.add_heading("PDF to Word Conversion", level=0)
    pdf_document = fitz.open(pdf_path)

    try:
        if not pages:
            document.add_paragraph("No text could be extracted from the uploaded PDF.")

        for page in pages:
            page_number = page.get("page_number")
            page_label = page_number if page_number is not None else "Unknown"
            normalized_page_number = _normalized_page_number(page_number)
            document.add_heading(f"Page {page_label}", level=1)

            blocks = sorted(page.get("blocks", []), key=_bbox_sort_key)
            has_text = False

            for block in blocks:
                text = _normalize_text(block.get("text", ""))
                if not text:
                    continue

                document.add_paragraph(text)
                has_text = True

            if not has_text:
                document.add_paragraph("No text detected on this page.")

            if normalized_page_number is not None:
                _append_page_images(
                    document,
                    pdf_document,
                    normalized_page_number,
                    images_by_page.get(normalized_page_number, [])
                )
    finally:
        pdf_document.close()

    _append_tables(document, tables)

    document.save(output_doc_path)
    return output_doc_path
